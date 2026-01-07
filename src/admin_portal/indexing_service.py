"""
Indexing Service for Practice Brain Admin Portal

Handles file processing, text extraction, chunking, and Pinecone indexing.
Uses the same logic as embed_data.py for consistency.
"""

import os
import uuid
import hashlib
import logging
from datetime import datetime
from typing import List, Dict, Optional, Tuple
from io import BytesIO

from langchain_openai import OpenAIEmbeddings
from langchain_experimental.text_splitter import SemanticChunker
from langchain_pinecone import Pinecone as LangchainPinecone
from langchain_core.documents import Document as LangchainDocument
import pinecone

from src.core.config import PINECONE_API_KEY, OPENAI_API_KEY, PINECONE_INDEX_NAME
from src.core.db import SessionLocal
from src.models.models import Document as DBDocument

logger = logging.getLogger(__name__)

# =============================================================================
# Configuration
# =============================================================================

# Chunk size configuration
MIN_CHUNK_SIZE = 100  # Minimum characters per chunk
MAX_CHUNK_SIZE = 2000  # Maximum characters per chunk

# Supported file types
SUPPORTED_EXTENSIONS = {
    ".txt": "text/plain",
    ".pdf": "application/pdf",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".doc": "application/msword",
    ".md": "text/markdown",
    ".html": "text/html",
    ".json": "application/json",
}


# =============================================================================
# Text Extraction Functions
# =============================================================================

def extract_text_from_txt(content: bytes) -> str:
    """Extract text from plain text file."""
    try:
        return content.decode('utf-8')
    except UnicodeDecodeError:
        return content.decode('latin-1')


def extract_text_from_pdf(content: bytes) -> str:
    """Extract text from PDF file using PyPDF2."""
    try:
        import PyPDF2
        pdf_reader = PyPDF2.PdfReader(BytesIO(content))
        text_parts = []
        for page in pdf_reader.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)
        return "\n\n".join(text_parts)
    except ImportError:
        logger.warning("PyPDF2 not installed. Trying pdfplumber...")
        try:
            import pdfplumber
            with pdfplumber.open(BytesIO(content)) as pdf:
                text_parts = []
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
            return "\n\n".join(text_parts)
        except ImportError:
            raise ImportError("Please install PyPDF2 or pdfplumber for PDF support: pip install PyPDF2 pdfplumber")
    except Exception as e:
        logger.error(f"Error extracting PDF text: {e}")
        raise


def extract_text_from_docx(content: bytes) -> str:
    """Extract text from DOCX file using python-docx."""
    try:
        from docx import Document as DocxDocument
        doc = DocxDocument(BytesIO(content))
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        return "\n\n".join(text_parts)
    except ImportError:
        raise ImportError("Please install python-docx for DOCX support: pip install python-docx")
    except Exception as e:
        logger.error(f"Error extracting DOCX text: {e}")
        raise


def extract_text_from_html(content: bytes) -> str:
    """Extract text from HTML file."""
    try:
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(content, 'html.parser')
        # Remove script and style elements
        for element in soup(['script', 'style', 'nav', 'footer', 'header']):
            element.decompose()
        text = soup.get_text(separator='\n')
        # Clean up whitespace
        lines = [line.strip() for line in text.splitlines() if line.strip()]
        return "\n".join(lines)
    except ImportError:
        raise ImportError("Please install beautifulsoup4 for HTML support: pip install beautifulsoup4")
    except Exception as e:
        logger.error(f"Error extracting HTML text: {e}")
        raise


def extract_text(filename: str, content: bytes) -> str:
    """
    Extract text from file based on extension.
    
    Args:
        filename: Original filename with extension
        content: File content as bytes
        
    Returns:
        Extracted text content
    """
    ext = os.path.splitext(filename.lower())[1]
    
    if ext == '.txt' or ext == '.md':
        return extract_text_from_txt(content)
    elif ext == '.pdf':
        return extract_text_from_pdf(content)
    elif ext in ['.docx', '.doc']:
        return extract_text_from_docx(content)
    elif ext == '.html':
        return extract_text_from_html(content)
    elif ext == '.json':
        return extract_text_from_txt(content)
    else:
        raise ValueError(f"Unsupported file type: {ext}")


# =============================================================================
# Chunking Functions
# =============================================================================

def simple_chunk_text(text: str, chunk_size: int = 800, overlap: int = 100) -> List[str]:
    """
    Simple text chunking with overlap.
    Fallback when semantic chunking is not available.
    """
    if len(text) <= chunk_size:
        return [text]
    
    chunks = []
    start = 0
    
    while start < len(text):
        end = start + chunk_size
        
        # Try to break at a sentence or paragraph boundary
        if end < len(text):
            # Look for paragraph break
            para_break = text.rfind('\n\n', start, end)
            if para_break > start + chunk_size // 2:
                end = para_break + 2
            else:
                # Look for sentence break
                for punct in ['. ', '! ', '? ', '.\n']:
                    sent_break = text.rfind(punct, start + chunk_size // 2, end)
                    if sent_break > 0:
                        end = sent_break + len(punct)
                        break
        
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        
        start = end - overlap
    
    return chunks


def semantic_chunk_text(text: str, embeddings: OpenAIEmbeddings) -> List[LangchainDocument]:
    """
    Semantic chunking using LangChain's SemanticChunker.
    Creates chunks based on semantic similarity.
    """
    try:
        text_splitter = SemanticChunker(embeddings)
        chunks = text_splitter.create_documents([text])
        return chunks
    except Exception as e:
        logger.warning(f"Semantic chunking failed, falling back to simple chunking: {e}")
        simple_chunks = simple_chunk_text(text)
        return [LangchainDocument(page_content=chunk) for chunk in simple_chunks]


# =============================================================================
# Indexing Service
# =============================================================================

class IndexingService:
    """Service for indexing documents into Pinecone."""
    
    def __init__(self):
        # Validate API keys
        if not OPENAI_API_KEY:
            raise ValueError("OPENAI_API_KEY environment variable is not set")
        if not PINECONE_API_KEY:
            raise ValueError("PINECONE_API_KEY environment variable is not set")
        if not PINECONE_INDEX_NAME:
            raise ValueError("PINECONE_INDEX_NAME environment variable is not set")
        
        logger.info(f"Initializing IndexingService with index: {PINECONE_INDEX_NAME}")
        logger.info(f"Using embedding model: text-embedding-3-small")
        
        self.embeddings = OpenAIEmbeddings(
            openai_api_key=OPENAI_API_KEY,
            model="text-embedding-3-small"
        )
        self.index_name = PINECONE_INDEX_NAME
        self.pc = pinecone.Pinecone(api_key=PINECONE_API_KEY)
    
    def compute_content_hash(self, content: str) -> str:
        """Compute hash of content for change detection."""
        return hashlib.sha256(content.encode('utf-8')).hexdigest()[:16]
    
    def process_and_index_file_with_progress(
        self,
        practice_id: str,
        filename: str,
        file_content: bytes,
        title: Optional[str] = None,
        source_type: str = "pdf",
        subagents_allowed: List[str] = None
    ):
        """
        Generator that yields progress updates during indexing.
        
        Yields:
            Dict with progress info: {"stage": str, "percent": int, "message": str}
        """
        doc_id = str(uuid.uuid4())
        title = title or os.path.splitext(filename)[0]
        subagents_allowed = subagents_allowed or ["chat", "clinical"]
        
        yield {"stage": "extracting", "percent": 5, "message": f"Extracting text from {filename}..."}
        
        try:
            # Step 1: Extract text (5-20%)
            text_content = extract_text(filename, file_content)
            
            if not text_content or len(text_content.strip()) < MIN_CHUNK_SIZE:
                yield {"stage": "error", "percent": 0, "message": "File contains insufficient text content", "error": True}
                return
            
            content_hash = self.compute_content_hash(text_content)
            char_count = len(text_content)
            yield {"stage": "extracted", "percent": 20, "message": f"Extracted {char_count:,} characters"}
            
            # Step 2: Chunk the text (20-50%)
            yield {"stage": "chunking", "percent": 25, "message": "Splitting text into semantic chunks..."}
            
            chunks = semantic_chunk_text(text_content, self.embeddings)
            
            if not chunks:
                yield {"stage": "error", "percent": 0, "message": "Failed to create chunks from content", "error": True}
                return
            
            chunk_count = len(chunks)
            yield {"stage": "chunked", "percent": 50, "message": f"Created {chunk_count} chunks"}
            
            # Step 3: Add metadata to chunks (50-60%)
            yield {"stage": "metadata", "percent": 55, "message": "Adding metadata to chunks..."}
            
            for i, chunk in enumerate(chunks):
                chunk.metadata = {
                    "doc_id": doc_id,
                    "practice_id": practice_id,
                    "source": filename,
                    "title": title,
                    "source_type": source_type,
                    "chunk_index": i,
                    "total_chunks": chunk_count,
                    "subagents_allowed": ",".join(subagents_allowed),
                    "indexed_at": datetime.utcnow().isoformat()
                }
            
            yield {"stage": "uploading", "percent": 60, "message": f"Uploading {chunk_count} chunks to Pinecone..."}
            
            # Step 4: Upload to Pinecone (60-95%)
            LangchainPinecone.from_documents(
                documents=chunks,
                embedding=self.embeddings,
                index_name=self.index_name,
                namespace=practice_id
            )
            
            yield {"stage": "uploaded", "percent": 95, "message": "Vectors uploaded to Pinecone"}
            
            # Done!
            yield {
                "stage": "complete",
                "percent": 100,
                "message": f"Successfully indexed '{title}' with {chunk_count} chunks",
                "result": {
                    "status": "success",
                    "doc_id": doc_id,
                    "title": title,
                    "chunk_count": chunk_count,
                    "content_hash": content_hash,
                    "text_length": char_count
                }
            }
            
        except ImportError as e:
            yield {"stage": "error", "percent": 0, "message": str(e), "error": True}
        except Exception as e:
            logger.error(f"Error indexing file: {e}")
            yield {"stage": "error", "percent": 0, "message": f"Indexing failed: {str(e)}", "error": True}
    
    def process_and_index_file(
        self,
        practice_id: str,
        filename: str,
        file_content: bytes,
        title: Optional[str] = None,
        source_type: str = "pdf",
        subagents_allowed: List[str] = None
    ) -> Dict:
        """
        Process a file and index it into Pinecone.
        
        Args:
            practice_id: The practice/client ID (used as Pinecone namespace)
            filename: Original filename
            file_content: Raw file content as bytes
            title: Document title (defaults to filename)
            source_type: Type of source (pdf, doc, etc.)
            subagents_allowed: Which agents can use this doc
            
        Returns:
            Dict with indexing results
        """
        doc_id = str(uuid.uuid4())
        title = title or os.path.splitext(filename)[0]
        subagents_allowed = subagents_allowed or ["chat", "clinical"]
        
        logger.info(f"Processing file: {filename} for practice: {practice_id}")
        
        try:
            # Step 1: Extract text
            logger.info(f"Extracting text from {filename}...")
            text_content = extract_text(filename, file_content)
            
            if not text_content or len(text_content.strip()) < MIN_CHUNK_SIZE:
                return {
                    "status": "error",
                    "message": "File contains insufficient text content",
                    "doc_id": doc_id
                }
            
            content_hash = self.compute_content_hash(text_content)
            logger.info(f"Extracted {len(text_content)} characters, hash: {content_hash}")
            
            # Step 2: Chunk the text
            logger.info("Chunking text semantically...")
            chunks = semantic_chunk_text(text_content, self.embeddings)
            
            if not chunks:
                return {
                    "status": "error",
                    "message": "Failed to create chunks from content",
                    "doc_id": doc_id
                }
            
            logger.info(f"Created {len(chunks)} chunks")
            
            # Step 3: Add metadata to chunks
            for i, chunk in enumerate(chunks):
                chunk.metadata = {
                    "doc_id": doc_id,
                    "practice_id": practice_id,
                    "source": filename,
                    "title": title,
                    "source_type": source_type,
                    "chunk_index": i,
                    "total_chunks": len(chunks),
                    "subagents_allowed": ",".join(subagents_allowed),
                    "indexed_at": datetime.utcnow().isoformat()
                }
            
            # Step 4: Upload to Pinecone
            logger.info(f"Uploading {len(chunks)} chunks to Pinecone namespace: {practice_id}")
            
            LangchainPinecone.from_documents(
                documents=chunks,
                embedding=self.embeddings,
                index_name=self.index_name,
                namespace=practice_id
            )
            
            logger.info(f"Successfully indexed document: {title}")

            # Persist document record to database
            try:
                db = SessionLocal()
                db_doc = DBDocument(
                    doc_id=uuid.UUID(doc_id),
                    client_id=uuid.UUID(practice_id),
                    title=title,
                    source_type=source_type,
                    source_uri=filename,
                    status="indexed",
                    chunk_count=len(chunks),
                    subagents_allowed=subagents_allowed,
                    last_indexed_at=datetime.utcnow()
                )
                db.add(db_doc)
                db.commit()
                db.close()
                logger.info(f"Document record saved to database: {doc_id}")
            except Exception as db_error:
                logger.error(f"Failed to save document to database: {db_error}")

            return {
                "status": "success",
                "message": f"Successfully indexed '{title}' with {len(chunks)} chunks",
                "doc_id": doc_id,
                "title": title,
                "chunk_count": len(chunks),
                "content_hash": content_hash,
                "text_length": len(text_content),
                "preview": text_content[:500] + "..." if len(text_content) > 500 else text_content
            }

        except ImportError as e:
            logger.error(f"Missing dependency: {e}")
            return {
                "status": "error",
                "message": str(e),
                "doc_id": doc_id
            }
        except Exception as e:
            logger.error(f"Error indexing file: {e}")
            return {
                "status": "error",
                "message": f"Indexing failed: {str(e)}",
                "doc_id": doc_id
            }
    
    def index_text_content(
        self,
        practice_id: str,
        title: str,
        text_content: str,
        source_type: str = "manual",
        source_uri: str = "manual-entry",
        subagents_allowed: List[str] = None
    ) -> Dict:
        """
        Index raw text content directly.
        
        Args:
            practice_id: The practice/client ID
            title: Document title
            text_content: Text to index
            source_type: Type of source
            source_uri: URI or description of source
            subagents_allowed: Which agents can use this doc
            
        Returns:
            Dict with indexing results
        """
        doc_id = str(uuid.uuid4())
        subagents_allowed = subagents_allowed or ["chat", "clinical"]
        
        try:
            if len(text_content.strip()) < MIN_CHUNK_SIZE:
                return {
                    "status": "error",
                    "message": "Text content is too short",
                    "doc_id": doc_id
                }
            
            content_hash = self.compute_content_hash(text_content)
            
            # Chunk the text
            chunks = semantic_chunk_text(text_content, self.embeddings)
            
            # Add metadata
            for i, chunk in enumerate(chunks):
                chunk.metadata = {
                    "doc_id": doc_id,
                    "practice_id": practice_id,
                    "source": source_uri,
                    "title": title,
                    "source_type": source_type,
                    "chunk_index": i,
                    "total_chunks": len(chunks),
                    "subagents_allowed": ",".join(subagents_allowed),
                    "indexed_at": datetime.utcnow().isoformat()
                }
            
            # Upload to Pinecone
            LangchainPinecone.from_documents(
                documents=chunks,
                embedding=self.embeddings,
                index_name=self.index_name,
                namespace=practice_id
            )

            # Persist document record to database
            try:
                db = SessionLocal()
                db_doc = DBDocument(
                    doc_id=uuid.UUID(doc_id),
                    client_id=uuid.UUID(practice_id),
                    title=title,
                    source_type=source_type,
                    source_uri=source_uri,
                    status="indexed",
                    chunk_count=len(chunks),
                    subagents_allowed=subagents_allowed,
                    last_indexed_at=datetime.utcnow()
                )
                db.add(db_doc)
                db.commit()
                db.close()
                logger.info(f"Document record saved to database: {doc_id}")
            except Exception as db_error:
                logger.error(f"Failed to save document to database: {db_error}")

            return {
                "status": "success",
                "message": f"Successfully indexed '{title}' with {len(chunks)} chunks",
                "doc_id": doc_id,
                "title": title,
                "chunk_count": len(chunks),
                "content_hash": content_hash
            }

        except Exception as e:
            logger.error(f"Error indexing text: {e}")
            return {
                "status": "error",
                "message": f"Indexing failed: {str(e)}",
                "doc_id": doc_id
            }

    def get_index_stats(self, practice_id: str) -> Dict:
        """Get statistics for a practice's namespace in Pinecone."""
        try:
            index = self.pc.Index(self.index_name)
            stats = index.describe_index_stats()
            
            namespaces = stats.get("namespaces", {})
            namespace_stats = namespaces.get(practice_id, {})
            
            return {
                "status": "success",
                "namespace": practice_id,
                "vector_count": namespace_stats.get("vector_count", 0),
                "total_vectors": stats.get("total_vector_count", 0)
            }
        except Exception as e:
            logger.error(f"Error getting index stats: {e}")
            return {
                "status": "error",
                "message": str(e)
            }
    
    def delete_document(self, practice_id: str, doc_id: str) -> Dict:
        """Delete all vectors for a specific document from Pinecone and database."""
        try:
            index = self.pc.Index(self.index_name)

            # Delete from Pinecone by metadata filter (doc_id)
            index.delete(
                namespace=practice_id,
                filter={"doc_id": {"$eq": doc_id}}
            )

            # Delete from database
            try:
                db = SessionLocal()
                db.query(DBDocument).filter(
                    DBDocument.doc_id == uuid.UUID(doc_id),
                    DBDocument.client_id == uuid.UUID(practice_id)
                ).delete()
                db.commit()
                db.close()
                logger.info(f"Document record deleted from database: {doc_id}")
            except Exception as db_error:
                logger.error(f"Failed to delete document from database: {db_error}")

            return {
                "status": "success",
                "message": f"Deleted document {doc_id} from namespace {practice_id}"
            }
        except Exception as e:
            logger.error(f"Error deleting document: {e}")
            return {
                "status": "error",
                "message": str(e)
            }


# Singleton instance
_indexing_service = None

def get_indexing_service() -> IndexingService:
    """Get or create the indexing service singleton."""
    global _indexing_service
    if _indexing_service is None:
        _indexing_service = IndexingService()
    return _indexing_service
